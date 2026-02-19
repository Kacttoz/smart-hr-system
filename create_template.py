from docx import Document

def create_rich_template():
    # Load original (renamed) template
    doc = Document('letters/template.docx')
    
    replacements = {
        '1010': '{{ reference_no }}',
        '12/February/2026': '{{ created_date }}',
        'Ms. Kirti Oberoi': '{{ candidate_name }}',
        'Kirti Oberoi': '{{ candidate_name }}',
        'HR Executive': '{{ designation }}',
        '16th February 2026': '{{ joining_date }}',
        '27,000/-': '{{ salary }}/-',
        'Twenty-Seven Thousand': '{{ salary_in_words }}',
    }

    def process_paragraph(paragraph):
        original_text = paragraph.text
        # Check if any replacement key is in the text
        clean_text = original_text
        found_keys = []
        
        # Identify which keys are present and their order
        # This is a simple implementation; complex overlapping keys might need better logic
        for key, value in replacements.items():
            if key in clean_text:
                found_keys.append((key, value))
        
        if not found_keys:
            return

        # Simple split and rebuild approach
        # Note: This removes original styling of the paragraph text (font, etc.) 
        # but preserves Paragraph styling (indentation, alignment).
        # We will make the 'value' runs BOLD.
        
        # Strategy: Replace keys with unique markers, then split by markers
        current_text = original_text
        for key, value in found_keys:
            # We use a unique marker to split safely
            marker = f"|||{key}|||"
            current_text = current_text.replace(key, marker)
            
        parts = current_text.split("|||")
        
        # Clear existing content
        paragraph.clear()
        
        for part in parts:
            # Check if this part corresponds to a key
            matched_key = None
            matched_value = None
            for key, value in found_keys:
                if part == key:
                    matched_key = key
                    matched_value = value
                    break
            
            if matched_key:
                # Add run with BOLD
                run = paragraph.add_run(matched_value)
                run.bold = True
                print(f"Bolding: {matched_value}")
            else:
                # Add normal text
                if part: # Skip empty strings
                    paragraph.add_run(part)

    # Process all paragraphs
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)
        
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)

    doc.save('letters/appointment_template.docx')
    print("Rich Template created at letters/appointment_template.docx")

if __name__ == "__main__":
    create_rich_template()
