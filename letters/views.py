from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from .models import OfferLetter
from django.http import HttpResponse
from django.conf import settings
from django.core.mail import EmailMessage
from docxtpl import DocxTemplate
import io
import os
from datetime import datetime

@login_required
def admin_dashboard(request):
    if not request.user.role == 'MAIN_ADMIN' and not request.user.is_superuser:
        return redirect('letters:sub_admin_dashboard')
        
    letters = OfferLetter.objects.all().order_by('-created_at')
    return render(request, 'letters/admin_dashboard.html', {'letters': letters})

@login_required
def sub_admin_dashboard(request):
    letters = OfferLetter.objects.filter(created_by=request.user).order_by('-created_at')
    return render(request, 'letters/sub_admin_dashboard.html', {'letters': letters})


from docxtpl import DocxTemplate
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.shared import Inches
import subprocess
import tempfile
import copy

def layout_corrections(docx_path, date_text, ref_suffix):
    """
    Ensures Reference Number and Date are on the same line (Ref Left, Date Right).
    Handles both separate paragraphs and same-paragraph (newline) cases.
    """
    doc = Document(docx_path)
    
    ref_paragraph = None
    date_paragraph = None
    
    def find_paragraphs(candidates):
        r_p = None
        d_p = None
        for p in candidates:
            if "KA/HR" in p.text or (str(ref_suffix) in p.text and "100" in str(ref_suffix)):
                 r_p = p
            if date_text in p.text:
                d_p = p
        return r_p, d_p

    # Collect all paragraphs from body and tables
    all_paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paragraphs.extend(cell.paragraphs)

    # Search linearly
    ref_paragraph, date_paragraph = find_paragraphs(all_paragraphs)
    
    if ref_paragraph and date_paragraph:
        # CASE 1: Same Paragraph (connected by newline likely)
        if ref_paragraph == date_paragraph:
            p = ref_paragraph
            # Replace newline/spaces between Ref and Date with a Tab
            # This is complex to do with Runs. Simplified approach:
            # Reconstruct the paragraph content.
            
            full_text = p.text
            # Identify Ref and Date parts
            # Assume Ref is at start, Date is at end? Or just search strings.
            if date_text in full_text:
                # Clear all runs
                for i in range(len(p.runs)):
                    p.runs[i].text = ""
                
                # Rebuild: Ref <TAB> Date
                # We need the Ref text (minus Date text if it was there)
                # Naive: clean text using replace
                clean_ref = full_text.replace(date_text, "").strip()
                
                run_ref = p.add_run(clean_ref)
                run_ref.font.bold = True
                
                p.add_run("\t")
                
                run_date = p.add_run(date_text)
                run_date.font.bold = True
                
                # Tab Stop (Increased to 7.0 inches to push date further right)
                p.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), WD_TAB_ALIGNMENT.RIGHT)

        # CASE 2: Different Paragraphs (Standard)
        else:
            # Move Date to Ref Paragraph
            # 1. Clear Date Paragraph
            try:
                p_element = date_paragraph._element
                p_element.getparent().remove(p_element)
            except:
                date_paragraph.clear() # Fallback

            # 2. Add Tab and Date to Ref Paragraph
            ref_paragraph.add_run("\t")
            r_d = ref_paragraph.add_run(date_text)
            r_d.font.bold = True
            
            # 3. Bold Ref Paragraph existing runs
            for run in ref_paragraph.runs:
                run.font.bold = True

            # 4. Tab Stop (Increased to 7.0 inches)
            ref_paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), WD_TAB_ALIGNMENT.RIGHT)
            
    doc.save(docx_path)

def bold_runs_in_docx(docx_path, texts_to_bold):
    """
    Opens a DOCX file, searches for specific texts, and bolds ONLY the specific text match,
    splitting runs if necessary.
    """
    doc = Document(docx_path)
    
    def process_paragraph(paragraph):
        for text_to_bold in texts_to_bold:
            if not text_to_bold or not text_to_bold.strip():
                continue
                
            for i, run in enumerate(paragraph.runs):
                if text_to_bold in run.text:
                    if run.text == text_to_bold:
                        run.font.bold = True
                    else:
                        try:
                            r_element = run._element
                            p_element = r_element.getparent()
                            parts = run.text.split(text_to_bold, 1)
                            pre_text = parts[0]
                            post_text = parts[1]
                            
                            run.text = pre_text 
                            
                            new_run_bold = copy.deepcopy(r_element)
                            new_run_bold.text = text_to_bold
                            rPr = new_run_bold.find(qn('w:rPr'))
                            if rPr is None:
                                rPr = OxmlElement('w:rPr')
                                new_run_bold.insert(0, rPr)
                            if rPr.find(qn('w:b')) is None:
                                OxmlElement('w:b')
                                rPr.append(OxmlElement('w:b'))
                            
                            p_element.insert(p_element.index(r_element) + 1, new_run_bold)
                            
                            if post_text:
                                new_run_post = copy.deepcopy(r_element)
                                new_run_post.text = post_text
                                p_element.insert(p_element.index(new_run_bold) + 1, new_run_post)
                        except:
                            pass

    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)
    doc.save(docx_path)

def generate_docx_file(letter):
    """Generates the DOCX file, applies corrections, and returns the path."""
    template_path = os.path.join(settings.BASE_DIR, 'letters', 'appointment_template.docx')
    doc = DocxTemplate(template_path)
    
    # Context values
    ctx_ref = f"{letter.id + 1050}"
    ctx_date = datetime.now().strftime("%d/%B/%Y")
    
    # Prefix Logic
    prefix = "Mr." if letter.gender == 'Male' else "Mrs."
    ctx_name = f"{prefix} {letter.candidate_name}"
    
    ctx_desg = letter.designation
    
    # Handle joining_date (might be string from POST immediate creation)
    if isinstance(letter.joining_date, str):
        try:
            date_obj = datetime.strptime(letter.joining_date, "%Y-%m-%d")
            ctx_join = date_obj.strftime("%d %B %Y")
        except ValueError:
            ctx_join = letter.joining_date
    else:
        ctx_join = letter.joining_date.strftime("%d %B %Y")
        
    try:
        ctx_sal = f"{float(letter.salary):.2f}"
    except (ValueError, TypeError):
        ctx_sal = str(letter.salary)
    
    # Sanitize salary words to remove "Rupees" and "Only" if they exist in DB (Legacy data)
    raw_words = letter.salary_in_words or ""
    clean_words = raw_words.lower().replace("rupees only", "").replace("rupees", "").replace("only", "").strip().title()
    ctx_sal_words = clean_words
    
    context = {
        'reference_no': ctx_ref,
        'created_date': ctx_date,
        'candidate_name': ctx_name,
        'designation': ctx_desg,
        'joining_date': ctx_join,
        'salary': ctx_sal,
        'salary_in_words': ctx_sal_words,
    }
    
    doc.render(context)
    
    temp_dir = tempfile.gettempdir()
    temp_docx_path = os.path.join(temp_dir, f"temp_letter_{letter.id}.docx")
    doc.save(temp_docx_path)
    
    # 1. Layout Corrections (Right Align Date, Bold Ref)
    try:
        layout_corrections(temp_docx_path, ctx_date, ctx_ref)
    except Exception as e:
        print(f"Layout correction failed: {e}")

    # 2. Bold Content Variables (Name, Salary, etc.)
    # Exclude Date and Ref (handled by layout_corrections)
    texts_to_bold = [ctx_name, ctx_desg, ctx_join, ctx_sal, ctx_sal_words]
    try:
        bold_runs_in_docx(temp_docx_path, texts_to_bold)
    except Exception as e:
        print(f"Bolding failed: {e}") 
    
    return temp_docx_path

@login_required
def generate_letter(request):
    if request.method == 'POST':
        candidate_name = request.POST.get('candidate_name')
        gender = request.POST.get('gender')
        email = request.POST.get('email')
        designation = request.POST.get('designation')
        salary = request.POST.get('salary')
        joining_date = request.POST.get('joining_date')
        
        # Determine User Action (Send or Preview)
        action = request.POST.get('action', 'send')
        
        # Instantiate without saving to DB yet
        letter = OfferLetter(
            candidate_name=candidate_name,
            gender=gender,
            email=email,
            designation=designation,
            salary=salary,
            joining_date=joining_date,
            created_by=request.user,
            status='DRAFT'
        )
        
        # Calculate Potential ID for Reference Number
        if action == 'preview':
            last_letter = OfferLetter.objects.last()
            next_id = (last_letter.id + 1) if last_letter else 1
            letter.id = next_id # Assign temporary ID for template rendering
        else:
            # For 'send', we MUST save to get a real ID and persist
            letter.save()

        try:
            # 1. Generate DOCX file
            docx_path = generate_docx_file(letter)
            
            # 2. Convert to PDF using LibreOffice
            pdf_path = docx_path.replace(".docx", ".pdf")
            out_dir = os.path.dirname(pdf_path)
            subprocess.run([
                "libreoffice", "--headless", "--convert-to", "pdf", 
                docx_path, "--outdir", out_dir
            ], check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            
            # 3. Handle Action
            if action == 'preview':
                # Return PDF directly
                with open(pdf_path, 'rb') as pdf_file:
                    pdf_content = pdf_file.read()
                    
                response = HttpResponse(pdf_content, content_type='application/pdf')
                response['Content-Disposition'] = f'inline; filename="Preview_{letter.candidate_name}.pdf"'
                
                # Cleanup temp files immediately
                if os.path.exists(docx_path): os.remove(docx_path)
                if os.path.exists(pdf_path): os.remove(pdf_path)
                
                return response

            elif action == 'send':
                # Send Email with PDF
                email_subject = f"Appointment Letter - {letter.designation}"
                
                # Prefix Logic for Email Body
                prefix = "Mr." if letter.gender == 'Male' else "Mrs."
                full_name = f"{prefix} {letter.candidate_name}"
                portal_id = str(letter.id + 1050)
                
                email_body = f"""Dear {full_name},

We are pleased to inform you that you have been selected for the position of {letter.designation} at Kactto.

Please find attached your official offer and joining kit detailing the terms and conditions of your work. We request you to go through the letter carefully and confirm your acceptance by replying to this email or signing and sending back a scanned copy of the letter within 48 hours from the date of issue. Else the letter will be declined by itself.


Portal id :- {portal_id}

Please find the login link below.

https://crmhr.kactto.com/login

If you have any questions or need further clarification, feel free to contact us my official email ID.
We look forward to welcoming you to our team.


Regards
HR Department
Hr@kactto.com"""
                
                email_msg = EmailMessage(
                    email_subject,
                    email_body,
                    settings.DEFAULT_FROM_EMAIL,
                    [email],
                )
                
                with open(pdf_path, 'rb') as f:
                    email_msg.attach(f"Appointment_Letter_{letter.candidate_name}.pdf", f.read(), 'application/pdf')
                
                email_msg.send(fail_silently=False)
                
                letter.status = 'SENT'
                letter.save()
                messages.success(request, f"Offer Letter generated and sent to {email} successfully!")
                
                # Cleanup temp files
                if os.path.exists(docx_path): os.remove(docx_path)
                if os.path.exists(pdf_path): os.remove(pdf_path)

        except Exception as e:
            # Handle Errors based on action
            if action == 'send':
                letter.status = 'FAILED'
                letter.save()
                messages.error(request, f"Letter generated but Email/Conversion failed: {str(e)}")
            else:
                 messages.error(request, f"Preview generation failed: {str(e)}")
                 # Ensure cleanup even on failure if paths exist
                 try:
                    if 'docx_path' in locals() and os.path.exists(docx_path): os.remove(docx_path)
                    if 'pdf_path' in locals() and os.path.exists(pdf_path): os.remove(pdf_path)
                 except: pass
        
        # Redirect back to dashboard (only if not returning preview response)
        return redirect('letters:sub_admin_dashboard')

    return render(request, 'letters/generate_form.html', {
        'designation_choices': OfferLetter.DESIGNATION_CHOICES
    })

@login_required
def download_doc(request, pk):
    """
    Downloads the letter as PDF.
    """
    letter = get_object_or_404(OfferLetter, pk=pk)
    # Check permission
    if not request.user.role == 'MAIN_ADMIN' and not request.user.is_superuser and letter.created_by != request.user:
        return HttpResponse("Unauthorized", status=403)
        
    try:
        # Generate DOCX
        docx_path = generate_docx_file(letter)
        
        # Convert to PDF using LibreOffice
        pdf_path = docx_path.replace(".docx", ".pdf")
        out_dir = os.path.dirname(pdf_path)
        subprocess.run([
            "libreoffice", "--headless", "--convert-to", "pdf", 
            docx_path, "--outdir", out_dir
        ], check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        
        # Read the PDF
        with open(pdf_path, 'rb') as pdf_file:
            response = HttpResponse(pdf_file.read(), content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="Appointment_Letter_{letter.candidate_name}.pdf"'
        
        # Cleanup
        if os.path.exists(docx_path): os.remove(docx_path)
        if os.path.exists(pdf_path): os.remove(pdf_path)
        
        return response

    except Exception as e:
        return HttpResponse(f"Error generating PDF: {str(e)}", status=500)
